# -*- coding: utf-8 -*-
r"""课标 Word 文档解析器（v5 全量稳定版）

输入：一份 .docx（全 Normal 样式，靠编号+关键词切分）
输出：条款记录列表 + JSON

层级识别（实测所有段落均为 Normal）：
  一级模块 : ^[一二三四五六七八九十]+、      例: 一、课程性质
  二级/主题: ^（[一二三四五六七八九十]+）     例: （一）物质
  条目     : ^\d+[.．、]                    例: 1.面向全体学生…
  内容标签 : 【内容要求】/【学业要求】/【教学提示】 （课程内容模块内）

两种文档结构（自动识别）：
  A. 常规型（义教/高中多数/课程方案）：正文用 一、二、三、 作一级模块，
     （一）（二） 作二级主题。前言(修订工作/指导思想)不在白名单→跳过。
  B. 括号型（如高中·数学）：正文没有 一、 一级标题，直接用
     （一）课程性质、（二）基本理念… 作一级模块；前言的 一、 仅出现在
     修订说明中（非白名单→跳过）。此时所有 （X）name（非目录/非噪声）
     都视为一级模块。

容错：docx 若含非法 XML 控制字符导致 python-docx 解析失败，自动修复后重试。
"""
import os, re, json, zipfile, tempfile, shutil
from collections import Counter

CN_NUM = "一二三四五六七八九十"
PUNCT = r"[.。．·․‧…]"
TOP_RE = re.compile(r"^([%s]+)[、.．]\s*(.+)$" % CN_NUM)
SUB_RE = re.compile(r"^（([%s]+)）\s*(.+)$" % CN_NUM)
ITEM_RE = re.compile(r"^(\d+)[.．、]\s*(.+)$")
CTYPE_RE = re.compile(
    r"^【?\s*(内容要求|学业要求|教学提示|教学建议|评价建议|教材编写建议|"
    r"课程资源开发与利用|教师培训建议|核心素养内涵|目标要求)\s*】?\s*[：:．.]*\s*$"
)

# 所有标准中出现过的一级模块名（义教/高中/课程方案 并集）
KNOWN_MODULES = {
    "课程性质", "课程性质与基本理念", "课程理念", "基本理念", "课程目标",
    "学科核心素养与课程目标", "课程内容", "课程结构", "学业质量", "课程实施",
    "实施建议", "附录",
    "培养目标", "基本原则", "课程设置", "课程标准编制与教材编写",
    "课程内容确定的原则", "课程实施与评价", "条件保障", "管理与监督",
    # 括号型文档（高中·数学）正文直接使用的模块名
    "学科核心素养", "设计依据", "结构", "学分与选课",
    "必修课程", "选择性必修课程", "选修课程",
    "学业质量内涵", "学业质量水平", "学业质量水平与考试评价的关系",
    "教学与评价建议", "学业水平考试与高考命题建议",
    "教材编写建议", "地方与学校实施课程标准的建议",
}
MODULE_TYPE = {
    "课程性质与基本理念": "课程性质", "课程性质": "课程性质",
    "课程理念": "课程理念", "基本理念": "课程理念",
    "课程目标": "课程目标", "学科核心素养与课程目标": "课程目标",
    "学科核心素养": "课程目标",
    "课程内容": "课程内容",
    "必修课程": "课程内容", "选择性必修课程": "课程内容", "选修课程": "课程内容",
    "课程结构": "课程结构", "设计依据": "课程结构", "结构": "课程结构",
    "学分与选课": "课程结构",
    "学业质量": "学业质量", "学业质量内涵": "学业质量",
    "学业质量水平": "学业质量", "学业质量水平与考试评价的关系": "学业质量",
    "课程实施": "课程实施", "实施建议": "课程实施",
    "教学与评价建议": "课程实施", "学业水平考试与高考命题建议": "课程实施",
    "教材编写建议": "课程实施", "地方与学校实施课程标准的建议": "课程实施",
    "附录": "附录",
    "培养目标": "培养目标", "基本原则": "基本原则", "课程设置": "课程设置",
    "课程标准编制与教材编写": "课程标准编制与教材编写",
    "课程内容确定的原则": "课程内容确定的原则",
    "课程实施与评价": "课程实施与评价",
    "条件保障": "条件保障", "管理与监督": "管理与监督",
}

# 目录行（含页码）识别：仅匹配真正的目录条目，避免误伤正文中的分数(1／2)等。
#   （一）课程性质 ／1          -> 括号编号 + 页码
#   课程内容…………………………33  -> 省略号引导页码（行尾）
TOC_LINE_RE = re.compile(r"[（(][^（）]{1,24}[）)]\s*[／/]\s*\d")
ELLIP_DIGIT_RE = re.compile(r"……\s*\d+\s*$")
# 页眉/标题噪声行（如“普通高中数学课程标准（2017年版2020年修订）”）：
# 仅匹配以 普通高中/义务教育/课程方案 开头且含“课程标准”的独立短行，避免误删正文中含“课程标准”的句子。
RUNNING_HEAD_RE = re.compile(r"^(普通高中|义务教育|课程方案)[^（(]{0,12}课程标准")
CJK_RE = re.compile(r"[一-鿿]")
# PDF→Word 转换残留的页脚碎片，如「■义务教育数学课程标准（2022年版）」，
# 常紧跟下一 clause 编号「（N）」。仅当含「年版」时清除，避免误删正文中
# 「（以下简称课程标准）」等合法引用。
FOOTER_RE = re.compile(
    r"■?\s*(?:义务教育|普通高中|课程方案)[^，。；：\n（(]{0,12}?课程标准\s*"
    r"[（(]\s*\d{4}\s*年[^）)]{0,6}[）)]\s*"
)


def _repair_xml_text(data):
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", data)


def load_paras(path):
    """返回段落文本列表；遇到非法控制字符自动修复后重试。"""
    import docx
    try:
        return [p.text for p in docx.Document(path).paragraphs]
    except Exception:
        pass
    tmp = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(path) as z:
            z.extractall(tmp)
        docp = os.path.join(tmp, "word", "document.xml")
        if os.path.exists(docp):
            data = open(docp, "rb").read().decode("utf-8", "ignore")
            data = _repair_xml_text(data)
            open(docp, "w", encoding="utf-8").write(data)
        fixed = os.path.join(tmp, "fixed.docx")
        if os.path.exists(fixed):
            os.remove(fixed)
        with zipfile.ZipFile(fixed, "w", zipfile.ZIP_DEFLATED) as z:
            for root, _, files in os.walk(tmp):
                for f in files:
                    p = os.path.join(root, f)
                    z.write(p, os.path.relpath(p, tmp))
        return [p.text for p in docx.Document(fixed).paragraphs]
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def clean_name(s):
    s = s.strip()
    s = re.sub(PUNCT + r"+\s*\d+\s*$", "", s)
    s = re.sub(r"\s+\d+\s*$", "", s)
    s = re.sub(PUNCT + r"+$", "", s)
    return s.strip()


def is_toc_entry(raw_name):
    """带页码/斜杠/省略号/次级括号的目录合并条目应跳过。
    例：「（一）物质 （二）运动和相互作用 14 …」「（一）核心素养内涵（二）目标要求」
    """
    s = raw_name.strip()
    if re.search(r"\d", s):
        return True
    if "／" in s or "/" in s:
        return True
    if "…" in s or "……" in s:
        return True
    if "（" in s or "(" in s:   # 合并目录项含次级括号
        return True
    return False


def detect_version(texts):
    for t in texts:
        m = re.search(r"[（(]\s*(\d{4})\s*年\s*版", t)
        if m:
            return m.group(1)
    return ""


def _prescan(paras):
    """返回 (uses_paren, body_start)。
    uses_paren=True 表示正文用 （X）name 作一级模块（无 一、 一级标题）。
    body_start 为应开始解析的段落下标（跳过标题/前言/目录）。
    """
    has_top = False
    last_toc = -1
    for i, raw in enumerate(paras):
        s = raw.strip()
        if not s:
            continue
        m = TOP_RE.match(s)
        if m and clean_name(m.group(2)) in KNOWN_MODULES:
            has_top = True
        if TOC_LINE_RE.search(s) or ELLIP_DIGIT_RE.search(s):
            last_toc = i
    uses_paren = not has_top
    body_start = last_toc + 1 if last_toc >= 0 else 0
    return uses_paren, body_start


def parse_paras(paras, xueduan, subject):
    version = detect_version(paras)
    uses_paren, _bs = _prescan(paras)
    # body_start 仅在「括号型」文档启用（用于跳过标题/前言/目录）；
    # 常规文档无需跳过——其前言 一、 不在白名单，已由白名单逻辑过滤。
    body_start = _bs if uses_paren else 0
    records = []
    started = False
    module = sub = ctype = item_no = None
    buf = []
    buf_title = ""

    def flush():
        nonlocal buf, buf_title, module, sub, ctype, item_no
        if not buf:
            return
        text = "\n".join(x.strip() for x in buf if x.strip())
        text = FOOTER_RE.sub("", text).strip()
        if not text.strip():
            return
        typ = MODULE_TYPE.get(module, module or "")
        records.append({
            "学段": xueduan, "学科": subject, "版本": version,
            "一级模块": module or "", "二级主题": sub or "",
            "条目编号": item_no or "", "条目标题": buf_title,
            "类型": typ, "原文": text,
        })
        buf = []

    for idx, raw in enumerate(paras):
        if idx < body_start:
            continue
        t = raw.strip()
        if not t:
            continue
        if not CJK_RE.search(t):          # 纯符号噪声行（如 $% &'()*+,-.）
            continue
        if RUNNING_HEAD_RE.search(t):      # 页眉/标题行
            continue
        if TOC_LINE_RE.search(t) or ELLIP_DIGIT_RE.search(t):  # 目录行
            continue

        m_top = TOP_RE.match(t)
        m_sub = SUB_RE.match(t)
        m_item = ITEM_RE.match(t)
        m_ctype = CTYPE_RE.match(t)

        if m_top:
            raw_name = m_top.group(2)
            name = clean_name(raw_name)
            if name not in KNOWN_MODULES or is_toc_entry(raw_name):
                continue
            flush()
            started = True
            module = name
            sub = ctype = item_no = None
            buf_title = t
            buf = [t]
            continue
        if m_sub:
            raw_name = m_sub.group(2)
            name = clean_name(raw_name)
            if is_toc_entry(raw_name):   # 目录合并条目（两种结构通用）
                continue
            if uses_paren:
                # 括号型：所有 （X）name 均为一级模块
                flush()
                started = True
                module = name
                sub = ctype = item_no = None
                buf_title = t
                buf = [t]
                continue
            else:
                # 常规型：（X）为二级主题，必须在一级模块内
                if not started:
                    continue
                flush()
                sub = name
                ctype = item_no = None
                buf_title = t
                buf = [t]
                continue
        if m_ctype:
            if not started:
                continue
            flush()
            ctype = m_ctype.group(1).strip()
            item_no = None
            buf_title = t
            buf = [t]
            continue
        if m_item and module != "课程内容":
            if not started:
                continue
            flush()
            item_no = m_item.group(1)
            buf_title = t
            buf = [t]
            continue
        if not started:          # 首个模块之前的标题/前言/目录噪声行，直接丢弃
            continue
        buf.append(t)
    flush()
    return records, version


def parse_docx(path, xueduan, subject):
    return parse_paras(load_paras(path), xueduan, subject)


def extract_subject(fn, cat):
    name = fn[:-5] if fn.lower().endswith(".docx") else fn
    name = re.sub(r"_[WP]0\d+$", "", name)
    if cat == "课程方案":
        return "普通高中课程方案" if ("普通高中" in name or "高中" in name) else "义务教育课程方案"
    return name


def run_all(base=None):
    if base is None:
        base = os.path.dirname(os.path.abspath(__file__))
    cats = {"义务教育": "义务教育", "高中": "普通高中", "课程方案": "课程方案"}
    all_recs, summary = [], []
    for cat, xueduan in cats.items():
        d = os.path.join(base, cat)
        if not os.path.isdir(d):
            continue
        for fn in sorted(os.listdir(d)):
            if not fn.lower().endswith(".docx"):
                continue
            subject = extract_subject(fn, cat)
            try:
                recs, ver = parse_docx(os.path.join(d, fn), xueduan, subject)
            except Exception as e:
                print("解析失败", fn, e)
                continue
            all_recs.extend(recs)
            summary.append((cat, subject, ver, len(recs)))
            with open(os.path.join(base, f"课标条款_{subject}_{xueduan}.json"),
                      "w", encoding="utf-8") as f:
                json.dump(recs, f, ensure_ascii=False, indent=2)
    total = os.path.join(base, "课标条款库_全20份.json")
    with open(total, "w", encoding="utf-8") as f:
        json.dump(all_recs, f, ensure_ascii=False, indent=2)
    print("=== 全量汇总 ===")
    for s in summary:
        print(f"  {s[0]:6s} {s[1]:8s} v{s[2]}  {s[3]}条")
    print(f"总计: {len(all_recs)} 条 -> {total}")
    return all_recs


def main():
    base = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base, "义务教育", "物理_W020220420582357585169.docx")
    recs, ver = parse_docx(path, "义务教育", "物理")
    out = os.path.join(base, "课标条款_物理_义教_试点.json")
    with open(out, "w", encoding="utf-8") as f:
        json.dump(recs, f, ensure_ascii=False, indent=2)
    print(f"版本: {ver}  条款数: {len(recs)}")
    print("一级模块:", dict(Counter(r["一级模块"] for r in recs)))
    print("类型分布:", dict(Counter(r["类型"] for r in recs)))


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--all":
        run_all()
    else:
        main()
